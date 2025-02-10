import json
import openai
import tkinter as tk
from tkinter import filedialog, Text, messagebox, font
import re
import textwrap
import sys
import os
from docx import Document

class TextEditorApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Text Editor App")

        default_model = "gpt-4o-mini"
        default_api_url = "https://your_api_url/v1/"

        self.model_var = tk.StringVar(value=default_model)
        self.api_url_var = tk.StringVar(value=default_api_url)

        # 设置字体
        self.custom_font = font.Font(family="微软雅黑", size=13)

        # 创建输入框
        input_frame = tk.Frame(root)
        input_frame.pack(side="top", fill="x", padx=10, pady=5)

        model_label = tk.Label(input_frame, text="模型名称：", font=self.custom_font)
        model_label.pack(side="left", padx=5)
        model_entry = tk.Entry(input_frame, textvariable=self.model_var, width=15, font=self.custom_font)
        model_entry.pack(side="left", padx=5)

        api_label = tk.Label(input_frame, text="API地址：", font=self.custom_font)
        api_label.pack(side="left", padx=5)
        api_entry = tk.Entry(input_frame, textvariable=self.api_url_var, width=20, font=self.custom_font)
        api_entry.pack(side="left", padx=5)

        process_button = tk.Button(input_frame, text="云处理", command=self.process_text, font=self.custom_font)
        process_button.pack(side="left", padx=10)

        frame = tk.Frame(root)
        frame.pack(expand='yes', fill='both')

        suggestion_frame = tk.Frame(frame, width=400, height=900)
        suggestion_frame.pack(expand=True,side='left', fill='both', padx=5, pady=5)
        suggestion_frame.pack_propagate(False)

        suggestion_title = tk.Label(suggestion_frame, text="修改建议", font=self.custom_font)
        suggestion_title.pack(side="top", anchor="w", padx=5, pady=5)

        self.suggestion_area = Text(suggestion_frame, wrap='char', font=self.custom_font)
        self.suggestion_area.pack(expand=True, fill='both')

        left_frame = tk.Frame(frame, width=700, height=900)
        left_frame.pack(expand=True,side='left', fill='both')
        left_frame.pack_propagate(False)

        left_title = tk.Label(left_frame, text="输入待处理的文本", font=self.custom_font)
        left_title.pack(side="top", anchor="w", padx=5, pady=5)

        self.text_area = Text(left_frame, wrap='char', font=self.custom_font)
        self.text_area.pack(expand=True, fill='both')

        right_frame = tk.Frame(frame, width=700, height=900)
        right_frame.pack(expand=True,side='right', fill='both')
        right_frame.pack_propagate(False)

        right_title = tk.Label(right_frame, text="完整的修改结果", font=self.custom_font)
        right_title.pack(side="top", anchor="w", padx=5, pady=5)

        self.modified_text_area = Text(right_frame, wrap='char', state='disabled', bg='#f0f0f0', font=self.custom_font)
        self.modified_text_area.pack(expand=True, fill='both')

        self.menu = tk.Menu(root)
        root.config(menu=self.menu)

        file_menu = tk.Menu(self.menu, tearoff=0)
        self.menu.add_cascade(label="File", menu=file_menu)
        file_menu.add_command(label="Open", command=self.open_file)
        file_menu.add_command(label="Save", command=self.save_file)
        file_menu.add_separator()
        file_menu.add_command(label="Exit", command=root.quit)

        self.tooltip = None
        self.sentence_positions = []
        self.sentence_tags = {}

        self.load_api_key()

    def load_config(self):
        # 打包后获取临时资源路径
        if getattr(sys, 'frozen', False):
            base_path = sys._MEIPASS
        else:
            base_path = os.path.dirname(os.path.abspath(__file__))
        
        config_path = os.path.join(base_path, 'config.json')
        with open(config_path, 'r', encoding='utf-8') as config_file:
            return json.load(config_file)

    def load_api_key(self):
        try:
            #with open('config.json') as config_file:
                #config = json.load(config_file)
                config = self.load_config()
                self.api_key = config["OPENAI_API_KEY"]
                openai.api_key = self.api_key
        except FileNotFoundError:
            messagebox.showerror("Error", "config.json file not found.")
        except KeyError:
            messagebox.showerror("Error", "OPENAI_API_KEY not found in config.json.")
        except json.JSONDecodeError:
            messagebox.showerror("Error", "Error decoding config.json.")

    def open_file(self):
        file_path = filedialog.askopenfilename(filetypes=[("Text Files", "*.txt"), ("Word Documents", "*.docx"), ("All Files", "*.*")])
        if file_path:
            if file_path.endswith('.docx'):
                self.open_docx_file(file_path)
            else:
                self.open_text_file(file_path)

    def open_text_file(self, file_path):
        with open(file_path, 'r', encoding='utf-8') as file:
            content = file.read()
            self.text_area.delete(1.0, tk.END)
            self.text_area.insert(tk.END, content)

    def open_docx_file(self, file_path):
        document = Document(file_path)
        content = ''
        for para in document.paragraphs:
            content += para.text + '\n'
        self.text_area.delete(1.0, tk.END)
        self.text_area.insert(tk.END, content)

    def save_file(self):
        file_path = filedialog.asksaveasfilename(defaultextension=".txt", filetypes=[("Text Files", "*.txt"), ("Word Documents", "*.docx"), ("All Files", "*.*")])
        if file_path:
            if file_path.endswith('.docx'):
                self.save_docx_file(file_path)
            else:
                self.save_text_file(file_path)

    def save_text_file(self, file_path):
        with open(file_path, 'w', encoding='utf-8') as file:
            content = self.text_area.get(1.0, tk.END)
            file.write(content)

    def save_docx_file(self, file_path):
        document = Document()
        content = self.text_area.get(1.0, tk.END).split('\n')
        for line in content:
            document.add_paragraph(line)
        document.save(file_path)
    
    def self_remove_spaces(self,text):
        symbols = '。，？！#……%（）{}【】[]：；“”‘'
        # 构建正则表达式模式（处理符号前后空格/全角空格/制表符）
        pattern = re.compile(r'([ \t\u3000]+(?=[{symbols}]))|((?<=[{symbols}])[ \t\u3000]+)'.format(symbols=re.escape(symbols)))
        return pattern.sub('', text)

    def process_text(self):
        content = self.text_area.get(1.0, tk.END)
        if not content.strip():
            messagebox.showwarning("Warning", "The text area is empty.")
            return

        # 移除中文标点符号前后的空格
        content = self.self_remove_spaces(content)

        self.text_area.delete(1.0, tk.END)
        self.text_area.insert(tk.END, content)

        openai.base_url = self.api_url_var.get()
        model_name = self.model_var.get()

        try:
            system_prompt = textwrap.dedent("""
                    你是一名非常专业的文稿编辑。你负责对文本内容的错别字、语法错误、内容错误、表述在中国有政治风险、难以理解的描述方法等进行优化，不用对文本进行润色，如果某句话没有上述问题就不用修改，无需对原文进行创作润色，务必保证你修改后的结果和作者原意一致。结果的输出总是以“最终我认为完美的修改结果：”开始，且修改内容按照如下JSON格式输出，同时使输出的JSON格式正确符合代码规范。举例：
                    我给你的输入：
                        请优化如下文本：推动两岸关系和平发展，必须继续坚持“和平统一、一郭两制”方针，退进祖国和平统一。
                    你的输入应该按照这个格式给出，在correct_text给出修改后的全文，需要保持原段落换行，在每一个修改details里，罗列每一处修改点；其中sentence_origin呈现本句原文，sentence_corrected呈现修改后的整句，modified_fragment记录这句修改的每处修改的片段，explain里写你这么修改的原因，ori_frag表示被修改片段的原文，correct_frag记录你给出的这部分片段的修改结果。
                        最终我认为完美的修改结果：
                        {
                            "item": {
                                "correct_text" : "推动两岸关系和平发展，必须继续坚持“和平统一、一国两制”方针，推进祖国和平统一。",
                                "details": [
                                    {
                                        "sentence_origin": "必须继续艰持“和平统一、一郭两制”方针，",
                                        "sentence_corrected": "必须继续坚持“和平统一、一国两制”方针，",
                                        "modified_fragment": [
                                            {
                                                "explain": "建议用“发展”替换“法展”，原因是“法”在这里是错别字。",
                                                "ori_frag": "法展",
                                                "correct_frag": "发展"
                                            },
                                            {
                                                "explain": "建议用“一国两制”替换“一郭两制”，原因是错别字，“一国两制”是专有名词。",
                                                "ori_frag": "一郭两制",
                                                "correct_frag": "一国两制"
                                            },
                                        ],
                                    },
                                    {
                                        "sentence_origin": "退进祖国和平统一。",
                                        "sentence_corrected": "推进祖国和平统一。",
                                        "modified_fragment": [
                                            {
                                                "explain": "建议用“推进”替换“退进”，退在这里是错别字",
                                                "ori_frag": "退进",
                                                "correct_frag": "推进"
                                            }
                                        ]
                                    }
                                ]
                            }
                        }
                """).strip()      
            response = openai.chat.completions.create(
                model=model_name, # 指定自定义模型名称          
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": f"请优化如下文本：{content}"}
                ],
                max_tokens=4096,  # 单次请求最大token数量
                temperature=1.0,
                stream=False
            )
            print(response)
            
            result_text = response.choices[0].message.content.strip()

            self.display_suggestions(content, result_text)

        except Exception as e:
            messagebox.showerror("Error", f"An error occurred: {e}")

    def display_suggestions(self, original_text, result_text):
        # 提取修改结果的 JSON 部分
        json_start = result_text.find('{')
        json_end = result_text.rfind('}') + 1
        json_str = result_text[json_start:json_end]
        modified_data = json.loads(json_str, strict=False)

        # 显示修改后的全文
        self.modified_text_area.config(state='normal')
        self.modified_text_area.delete(1.0, tk.END)
        self.modified_text_area.insert(tk.END, modified_data['item']['correct_text'])
        self.modified_text_area.config(state='disabled')

        differences = modified_data['item']['details']
        if not differences:
            messagebox.showinfo("Info", "没有找到错误!")
            return

        self.sentence_positions = []  # 清空之前的句子位置记录
        self.sentence_tags = {}  # 清空之前的句子标签记录

        for idx, diff in enumerate(differences):
            sentence_origin = diff['sentence_origin']
            sentence_corrected = diff['sentence_corrected']
            modified_fragments = diff['modified_fragment']

            start_sentence_idx = self.text_area.search(sentence_origin, "1.0", tk.END)
            if start_sentence_idx:
                end_sentence_idx = f"{start_sentence_idx} + {len(sentence_origin)}c"
                self.sentence_positions.append((start_sentence_idx, end_sentence_idx))

                for frag_idx, fragment in enumerate(modified_fragments):
                    ori_frag = fragment['ori_frag']
                    correct_frag = fragment['correct_frag']
                    explain = fragment['explain']

                    start_idx = self.text_area.search(ori_frag, start_sentence_idx, end_sentence_idx)
                    if start_idx:
                        end_idx = f"{start_idx} + {len(ori_frag)}c"
                        tag_name_ori = f"suggestion_ori_{idx}_{frag_idx}"
                        tag_name_correct = f"suggestion_correct_{idx}_{frag_idx}"
                        self.text_area.tag_add(tag_name_ori, start_idx, end_idx)
                        self.text_area.tag_config(tag_name_ori, background="#f8d7da", foreground="#e25b68")
                        self.text_area.insert(end_idx, correct_frag)
                        self.text_area.tag_add(tag_name_correct, end_idx, f"{end_idx} + {len(correct_frag)}c")
                        self.text_area.tag_config(tag_name_correct, background="#d1e7dd", foreground="#248c5c")
                        self.text_area.tag_bind(tag_name_ori, "<Enter>", lambda e, s=correct_frag, o=ori_frag, expl=explain, ori=sentence_origin, corr=sentence_corrected, idx=start_idx: self.show_suggestion_tooltip(e, s, o, expl, ori, corr, idx))
                        self.text_area.tag_bind(tag_name_ori, "<Leave>", self.hide_suggestion_tooltip)
                        self.text_area.tag_bind(tag_name_ori, "<Button-1>", lambda e, s=correct_frag, o=ori_frag, expl=explain, ori=sentence_origin, corr=sentence_corrected, t_ori=tag_name_ori, t_corr=tag_name_correct: self.display_suggestion_in_area(s, o, expl, ori, corr, t_ori, t_corr))

                        if sentence_origin not in self.sentence_tags:
                            self.sentence_tags[sentence_origin] = []
                        self.sentence_tags[sentence_origin].append(tag_name_ori)

    def display_suggestion_in_area(self, suggestion, ori_frag, explanation, sentence_origin, sentence_corrected, tag_name_ori, tag_name_correct):
        self.suggestion_area.config(state='normal')
        self.suggestion_area.delete(1.0, tk.END)

        label = tk.Label(self.suggestion_area, text=f"修改建议：\n（1）将 {ori_frag} 修改为 {suggestion}；\n（2）原因是：{explanation}；\n（3）\n本句原文：{sentence_origin}\n本句改后：{sentence_corrected}", justify="left", wraplength=380, font=self.custom_font)
        self.suggestion_area.window_create(tk.END, window=label)
        self.suggestion_area.insert(tk.END, "\n\n")

        btn_replace_fragment = tk.Button(self.suggestion_area, text="仅替换片段", command=lambda: self.replace_fragment(tag_name_ori, tag_name_correct), font=self.custom_font)
        self.suggestion_area.window_create(tk.END, window=btn_replace_fragment)
        self.suggestion_area.insert(tk.END, " ")

        btn_replace_sentence = tk.Button(self.suggestion_area, text="替换整句", command=lambda: self.replace_sentence(tag_name_ori, tag_name_correct), font=self.custom_font)
        self.suggestion_area.window_create(tk.END, window=btn_replace_sentence)
        self.suggestion_area.insert(tk.END, " ")

        btn_cancel = tk.Button(self.suggestion_area, text="不做修改", command=lambda: self.remove_correct_frag(tag_name_correct), font=self.custom_font)
        self.suggestion_area.window_create(tk.END, window=btn_cancel)
        self.suggestion_area.insert(tk.END, " ")

        btn_close = tk.Button(self.suggestion_area, text="取消", command=lambda: self.clear_suggestion_area(), font=self.custom_font)
        self.suggestion_area.window_create(tk.END, window=btn_close)

        self.suggestion_area.config(state='disabled')

    def clear_suggestion_area(self):
        self.suggestion_area.config(state='normal')
        self.suggestion_area.delete(1.0, tk.END)
        self.suggestion_area.config(state='disabled')

    def replace_fragment(self, tag_name_ori, tag_name_correct):
        start_idx, end_idx = self.text_area.tag_ranges(tag_name_ori)
        self.text_area.delete(start_idx, end_idx)
        self.text_area.tag_delete(tag_name_ori)

    def replace_sentence(self, tag_name_ori, tag_name_correct):
        sentence_origin = ""
        for key, value in self.sentence_tags.items():
            if tag_name_ori in value:
                sentence_origin = key
                break

        if sentence_origin in self.sentence_tags:
            for tag_name in self.sentence_tags[sentence_origin]:
                ranges = self.text_area.tag_ranges(tag_name)
                if len(ranges) == 2:  # 确保 tag_ranges 返回有效的索引
                    start_idx, end_idx = ranges
                    self.text_area.delete(start_idx, end_idx)
                    self.text_area.tag_delete(tag_name)
            self.sentence_tags.pop(sentence_origin)

    def remove_correct_frag(self, tag_name_correct):
        start_idx, end_idx = self.text_area.tag_ranges(tag_name_correct)
        self.text_area.delete(start_idx, end_idx)
        self.text_area.tag_delete(tag_name_correct)

    def show_suggestion_tooltip(self, event, suggestion, ori_frag, explanation, sentence_origin, sentence_corrected, index):
        bbox = self.text_area.bbox(index)
        if bbox:
            x, y, width, height = bbox
            self.tooltip = tk.Toplevel(self.text_area)
            self.tooltip.wm_overrideredirect(True)
            self.tooltip.wm_geometry(f"+{x+self.text_area.winfo_rootx()}+{y+self.text_area.winfo_rooty() + height}")
            tooltip_text = f"修改建议：\n（1）将 {ori_frag} 修改为 {suggestion}；\n（2）原因是：{explanation}；\n（3）本句原文：{sentence_origin}\n本句改后：{sentence_corrected}"
            label = tk.Label(self.tooltip, text=tooltip_text, background="yellow", relief="solid", borderwidth=1, justify="left", font=self.custom_font)
            label.pack()

    def hide_suggestion_tooltip(self, event):
        if self.tooltip:
            self.tooltip.destroy()
            self.tooltip = None


if __name__ == "__main__":
    root = tk.Tk()
    app = TextEditorApp(root)
    root.mainloop()
